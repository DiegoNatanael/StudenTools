# doc_generator.py
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn # qn might not be strictly needed for basic doc generation

def create_document_docx(title_text: str, intro_text: str, items: list, table_data: list) -> bytes:
    """
    Generates a .docx document based on provided data and returns its bytes.
    """
    doc = Document()

    # --- Cover Page ---
    doc.add_heading(title_text, level=0)
    p = doc.add_paragraph('Generado automáticamente por StudenTools. ')
    p.runs[0].font.size = Pt(12)
    doc.add_paragraph() # blank line

    # --- Section: Introduction ---
    doc.add_heading('1. Introducción', level=1)
    doc.add_paragraph(intro_text)

    # --- Section: Content with styles and emphasis ---
    doc.add_heading('2. Contenido principal', level=1)
    p = doc.add_paragraph()
    run = p.add_run("Texto normal, seguido de ")
    run.font.size = Pt(11)
    run = p.add_run("texto en negrita")
    run.bold = True
    run.font.size = Pt(11)
    p.add_run(" y ")
    run2 = p.add_run("texto en cursiva.")
    run2.italic = True
    run2.font.size = Pt(11)

    # --- Bullet List ---
    doc.add_heading('3. Lista de puntos', level=2)
    for it in items:
        doc.add_paragraph(it, style='List Bullet')

    # --- Table ---
    doc.add_heading('4. Tabla de ejemplo', level=2)
    table = doc.add_table(rows=1, cols=len(table_data[0] if table_data else 3)) # Dynamically set cols
    hdr_cells = table.rows[0].cells
    # Assuming first row of table_data is headers, or provide explicit headers
    if table_data and isinstance(table_data[0], list) and len(table_data[0]) == 3: # Simple header assumption
        hdr_cells[0].text = 'Columna A'
        hdr_cells[1].text = 'Columna B'
        hdr_cells[2].text = 'Columna C'
        actual_rows = table_data
    else: # Fallback or if table_data provides headers explicitly
        hdr_cells[0].text = 'Header 1'
        hdr_cells[1].text = 'Header 2'
        hdr_cells[2].text = 'Header 3'
        actual_rows = table_data


    for row_data in actual_rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            if i < len(row_cells): # Ensure we don't go out of bounds if rows have fewer columns
                row_cells[i].text = str(val)

    table.autofit = True

    # --- Section: Conclusion ---
    doc.add_heading('5. Conclusión', level=1)
    doc.add_paragraph("Aquí puedes poner conclusiones, próximos pasos y referencias.")

    # --- Footer ---
    section = doc.sections[0]
    footer = section.footer
    f_p = footer.paragraphs[0]
    f_p.text = "Documento generado automáticamente • Página "
    f_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Save to a BytesIO object instead of a file
    from io import BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0) # Rewind to the beginning of the buffer
    return buffer.getvalue()
